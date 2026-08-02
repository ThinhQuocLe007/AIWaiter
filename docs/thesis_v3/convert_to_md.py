#!/usr/bin/env python3
"""
Convert thesis_v3.docx to multiple markdown files split by chapter (Heading 1).
Extracts embedded images and writes them to a images/ directory.
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

# ── paths ─────────────────────────────────────────────────────────────
DOCX_PATH = "docs/thesis_v3/thesis_v3.docx"
OUT_DIR = Path("docs/thesis_v3/markdown")
IMG_DIR = OUT_DIR / "images"

# ── parse images from docx ───────────────────────────────────────────
def extract_images(doc, img_dir: Path):
    """Extract all embedded images from the docx to img_dir."""
    img_dir.mkdir(parents=True, exist_ok=True)
    mapping = {}  # rId -> filename

    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            ext = os.path.splitext(rel.target_ref)[1]
            # consistent naming: image1.png, etc.
            try:
                num = int(re.search(r"image(\d+)", rel.target_ref).group(1))
            except (AttributeError, ValueError):
                num = hash(rel.target_ref) % 10000
            fname = f"image{num:02d}{ext}"
            mapping[rel_id] = fname
            filepath = img_dir / fname
            if not filepath.exists():
                filepath.write_bytes(rel.target_part.blob)
    return mapping

def find_images_in_paragraph(para):
    """
    Check if a paragraph contains inline drawings/images.
    Returns list of (rId, blip_embed) tuples.
    """
    images = []
    nsmap = {
        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    }
    for blip in para._element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
        embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
        if embed:
            images.append(embed)
    return images

# ── paragraph → markdown helpers ──────────────────────────────────────

def clean_text(text: str) -> str:
    """Clean text: strip, normalize whitespace, escape markdown special chars lightly."""
    return text.strip()

def style_to_md_level(style_name: str) -> int:
    if style_name == "Heading 1":
        return 1
    elif style_name == "Heading 2":
        return 2
    elif style_name == "Heading 3":
        return 3
    return 0

def para_to_md(para, img_map: dict, seen_images: set) -> str:
    """Convert a single paragraph to markdown text."""
    text = para.text
    style = para.style.name if para.style else "Normal"

    # Check for images in this paragraph
    image_embeds = find_images_in_paragraph(para)
    img_md = ""
    for embed in image_embeds:
        if embed in img_map and img_map[embed] not in seen_images:
            fname = img_map[embed]
            seen_images.add(fname)
            img_md += f"![{fname}](images/{fname})\n\n"

    # Heading
    level = style_to_md_level(style)
    if level > 0:
        clean = clean_text(text)
        if clean:
            return f"{img_md}{'#' * level} {clean}\n\n"
        return img_md

    # Figure paragraph
    if style == "Figure":
        clean = clean_text(text)
        if clean:
            return f"{img_md}*{clean}*\n\n"
        return img_md

    # Caption paragraph (tables usually)
    if style == "Caption":
        clean = clean_text(text)
        if clean:
            return f"{img_md}*{clean}*\n\n"
        return img_md

    # List paragraph
    if style == "List Paragraph":
        clean = clean_text(text)
        if clean:
            return f"{img_md}- {clean}\n"
        return img_md

    # Table of figures
    if style == "table of figures":
        clean = clean_text(text)
        if clean:
            return f"{img_md}{clean}\n"
        return img_md

    # No Spacing
    if style == "No Spacing":
        clean = clean_text(text)
        if clean:
            return f"{img_md}{clean}\n\n"
        return img_md

    # Normal text
    clean = clean_text(text)
    if clean:
        return f"{img_md}{clean}\n\n"
    return img_md if img_md else "\n"

def convert_table_to_md(table) -> str:
    """Convert a docx table to markdown table."""
    rows = table.rows
    if not rows:
        return ""

    lines = []
    max_cols = 0

    # First pass: determine column count
    for row in rows:
        max_cols = max(max_cols, len(row.cells))

    if max_cols == 0:
        return ""

    for ri, row in enumerate(rows):
        cells = []
        col_idx = 0
        for cell in row.cells:
            text = cell.text.strip().replace("\n", " ")
            # Handle merged cells (grid_span / vmerge may repeat content)
            cells.append(text)
        # Pad short rows
        while len(cells) < max_cols:
            cells.append("")

        lines.append("| " + " | ".join(cells) + " |")

        # Header separator after first row
        if ri == 0:
            lines.append("| " + " | ".join(["---"] * max_cols) + " |")

        # Only take first row for header-like tables with duplicate rows
        # (the thesis has many 1-row description tables)

    return "\n".join(lines) + "\n\n"

# ── chapter detection ─────────────────────────────────────────────────

CHAPTER_PATTERNS = [
    ("CHAPTER 1", "01_chapter_1_introduction"),
    ("CHAPTER 2", "02_chapter_2_related_work"),
    ("CHAPTER 3", "03_chapter_3_robot_navigation"),
    ("CHAPTER 4", "04_chapter_4_ai_backend"),
    ("CHAPTER 5", "05_chapter_5_experimental_results"),
    ("CHAPTER 6", "06_chapter_6_conclusion"),
]

FRONT_MATTER_TITLES = [
    "STATEMENT OF COMMITMENT",
    "ACKNOWLEDGMENT",
    "ABSTRACT",
    "LIST OF FIGURES",
    "LIST OF TABLES",
    "LIST OF ABBREVIATION",
]

def detect_chapter(text: str) -> str | None:
    """Return a chapter slug from a heading 1 text, or None for front matter."""
    upper = text.upper().replace("\n", " ")

    for pattern, slug in CHAPTER_PATTERNS:
        if pattern in upper:
            return slug

    if "REFERENCES" in upper:
        return "07_references"

    if "APPENDIX" in upper:
        return "08_appendix"

    for fm in FRONT_MATTER_TITLES:
        if fm.upper() in upper:
            return "00_front_matter"

    return "00_front_matter"

# ── main conversion ───────────────────────────────────────────────────

def main():
    doc = Document(DOCX_PATH)
    img_map = extract_images(doc, IMG_DIR)
    print(f"Extracted {len(img_map)} images to {IMG_DIR}")

    OUT_DIR.mkdir(parents=True, exist_ok=True)

    # Build ordered list of elements: paragraphs, tables
    # We need to track where tables appear relative to paragraphs
    # python-docx order: paragraphs[0..n], tables[0..m], in document.xml order

    body = doc.element.body
    para_elements = list(body.iterchildren(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'
    ))
    table_elements = list(body.iterchildren(
        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl'
    ))

    # Build the element order
    elem_order = []
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            elem_order.append(('p', child))
        elif tag == 'tbl':
            elem_order.append(('tbl', child))

    para_idx_map = {}
    for i, p in enumerate(para_elements):
        para_idx_map[id(p)] = i

    # We'll collect content into chapter buckets
    chapters = {}  # slug -> list of (kind, content)
    current_chapter = "00_front_matter"
    chapters[current_chapter] = []
    seen_images = set()

    # Also track what the current section title is for front matter
    current_front_section = None

    para_index = 0
    for kind, elem in elem_order:
        if kind == 'p':
            para = doc.paragraphs[para_index]
            style = para.style.name if para.style else "Normal"
            text = para.text

            para_index += 1

            # Check for CHAPTER 4 heading embedded in P669
            if "CHAPTER 4" in text and style != "Heading 1":
                # Split: part before "CHAPTER 4" stays in current chapter
                # "CHAPTER 4: ..." starts a new chapter
                split_idx = text.upper().find("CHAPTER 4")
                if split_idx > 0:
                    before = text[:split_idx].strip()
                    chapter_text = text[split_idx:].strip()

                    if before:
                        chapters[current_chapter].append(('md', para_to_md(para, img_map, seen_images)))

                    # Start Chapter 4
                    for p, slug in CHAPTER_PATTERNS:
                        if p in chapter_text.upper():
                            current_chapter = slug
                            chapters[current_chapter] = []
                            chapters[current_chapter].append(('md', f"# {chapter_text}\n\n"))
                            break
                    continue
                elif text.upper().startswith("CHAPTER 4"):
                    for p, slug in CHAPTER_PATTERNS:
                        if p in text.upper():
                            current_chapter = slug
                            chapters[current_chapter] = []
                            chapters[current_chapter].append(('md', f"# {clean_text(text)}\n\n"))
                            break
                    continue

            # Heading 1 → new chapter
            if style == "Heading 1":
                clean = clean_text(text)
                slug = detect_chapter(clean)
                if slug != current_chapter:
                    current_chapter = slug
                    if current_chapter not in chapters:
                        chapters[current_chapter] = []
                    if clean:
                        chapters[current_chapter].append(('md', f"# {clean}\n\n"))
                else:
                    # Same chapter (front matter has multiple Heading 1)
                    if clean:
                        chapters[current_chapter].append(('md', f"# {clean}\n\n"))
                continue

            # Normal paragraph
            md = para_to_md(para, img_map, seen_images)
            if md:
                chapters[current_chapter].append(('md', md))

        elif kind == 'tbl':
            # Find table index
            tbl_idx = table_elements.index(elem)
            table = doc.tables[tbl_idx]
            md = convert_table_to_md(table)
            if md:
                chapters[current_chapter].append(('md', md))

    # ── Write chapter files ───────────────────────────────────────────
    for slug, content in chapters.items():
        if not content:
            continue

        # Build filename
        if slug == "00_front_matter":
            fname = "00_front_matter.md"
        else:
            fname = f"{slug}.md"

        # Build content
        lines = []
        for kind, txt in content:
            lines.append(txt)

        full_text = "".join(lines)

        # Clean up: collapse multiple blank lines
        full_text = re.sub(r'\n{4,}', '\n\n\n', full_text)

        filepath = OUT_DIR / fname
        filepath.write_text(full_text, encoding="utf-8")
        print(f"  Wrote {fname}  ({len(full_text)} chars)")

    print(f"\nDone. {len(chapters)} files written to {OUT_DIR}")

if __name__ == "__main__":
    main()
