#!/usr/bin/env python3
"""
Convert thesis.docx to markdown files split by chapter.
Output: docs/thesis_final_version/md/
"""

import re
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

SRC = Path(__file__).parent / "thesis.docx"
OUT_DIR = Path(__file__).parent / "md"
OUT_DIR.mkdir(exist_ok=True)

doc = Document(SRC)


# ---- Markdown helpers ----
def run_to_md(run):
    """Convert a single run (text span) to markdown with formatting."""
    text = run.text
    if not text:
        return ""
    if run.bold and run.italic:
        return f"***{text}***"
    if run.bold:
        return f"**{text}**"
    if run.italic:
        return f"*{text}*"
    return text


def para_to_md(para, list_counter=None):
    """Convert a paragraph to markdown."""
    style = para.style.name if para.style else "Normal"
    text = "".join(run_to_md(r) for r in para.runs).strip()
    if not text:
        return ""

    # Heading
    if style.startswith("Heading"):
        level = int(style.split()[-1])
        return "\n" + "#" * level + " " + text + "\n"

    # List paragraph
    if style == "List Paragraph" or style == "List Bullet" or style == "List Number":
        return "- " + text + "\n"

    return text + "\n\n"


def extract_table_as_md(table):
    """Convert a docx table to GitHub-flavored markdown table."""
    rows = table.rows
    if not rows:
        return ""

    md_rows = []
    for row in rows:
        cells = []
        for cell in row.cells:
            cell_text = " ".join(
                p.text.strip() for p in cell.paragraphs if p.text.strip()
            )
            cell_text = cell_text.replace("\n", " ").replace("|", "\\|")
            cells.append(cell_text)
        md_rows.append("| " + " | ".join(cells) + " |")

    if not md_rows:
        return ""

    # GFM table needs header separator
    header = md_rows[0]
    num_cols = len(header.split("|")) - 2
    sep = "|" + "|".join([" --- " for _ in range(num_cols)]) + "|"

    return "\n" + "\n".join([header, sep] + md_rows[1:]) + "\n\n"


# ---- Build element index ----
# docx paragraphs and tables are separate; we need to interleave them by position.
class Element:
    def __init__(self, idx, etype, obj):
        self.idx = idx
        self.etype = etype  # "para" or "table"
        self.obj = obj


# Get paragraph positions via XML
body = doc.element.body
elements = []
para_idx = 0
table_idx = 0

for child in body:
    tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
    if tag == "p":
        elements.append(Element(para_idx, "para", doc.paragraphs[para_idx]))
        para_idx += 1
    elif tag == "tbl":
        elements.append(Element(table_idx, "table", doc.tables[table_idx]))
        table_idx += 1

# ---- Split by Heading 1 chapters ----
chapters = {}  # title -> list of markdown strings
current_title = "FRONT_MATTER"
chapters[current_title] = []

for el in elements:
    if el.etype == "para":
        para = el.obj
        style = para.style.name if para.style else "Normal"
        text = para.text.strip()

        # Detect Heading 1
        if style == "Heading 1":
            current_title = text
            if current_title not in chapters:
                chapters[current_title] = []
            chapters[current_title].append("# " + text + "\n\n")
        else:
            md = para_to_md(para)
            if md:
                chapters[current_title].append(md)

    elif el.etype == "table":
        table = el.obj
        md = extract_table_as_md(table)
        if md:
            chapters[current_title].append(md)


# ---- Normalize chapter filenames ----
def slugify(title):
    """Convert chapter title to filename."""
    name = title.lower()
    name = re.sub(r"[^a-z0-9\s-]", "", name)
    name = re.sub(r"\s+", "_", name.strip())
    return name


chapter_order = [k for k in chapters.keys()]

for i, title in enumerate(chapter_order):
    slug = slugify(title)
    # Number prefix for correct ordering
    prefix = f"{i:02d}_"
    filename = f"{prefix}{slug}.md"
    content = "".join(chapters[title])
    out_path = OUT_DIR / filename
    out_path.write_text(content, encoding="utf-8")
    print(f"  [{len(content):>6} chars] {filename}")

print(f"\nDone! {len(chapter_order)} files written to {OUT_DIR}")
